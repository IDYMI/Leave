import datetime
import time

from docx import Document

import bin as bin
from lib.db.Cur import get_cur as get_cur, end_conn as end_conn
from lib.docx.docx_add import add_info as add_info
from lib.docx.docx_add import add_text as add_text
from lib.docx.get_Times import get_Times as get_Times
from lib.docx.replace_date import replace_date_in_docx as replace_date_in_docx

title_text = '请假条'
race_text = '备战蓝桥杯，ICPC'
race_text = '参加暑假集训选拔赛'
date_range = '01月01日 - 01月02日'

traning = "1"

year = datetime.datetime.now().year
end_date = f'{year}年01月01日'
end_text = '计算机与软件学院计科ACM集训队'
normal_size = 12
StartTime = time.strftime("%m{m}%d{d}", time.localtime()).format(m='月', d='日')  # 当前的时间
StartTime = '6月14日'
Pause = 3  # 一次时长
T = 1  # 连续几次
Times = get_Times(StartTime, Times=T, Pause=Pause)

if __name__ == '__main__':
    conn, cur = get_cur()

    doc = Document()  # 新建一个空白页

    Class = ['20', '21', '22']
    Class = ['22']

    # Class = str(tuple(Class)).replace(",)", ")")

    for idx, c in enumerate(Class):
        """遍历届数"""
        for idx2, Time in enumerate(Times):
            # print(Time)
            """遍历时间域"""

            # 插入标题
            title = doc.add_heading()
            add_text(title, title_text, 26, Bold=True, alignment='CENTER')

            # 插入原因
            reason = doc.add_paragraph()

            if traning != '全天':
                if int(str(year)[2:]) - int(c) > 2:  # 大三
                    place = '13#612'
                else:  # 大一大二
                    place = '13#602'
                traning = f"早7：00-晚10：00在{place}训练"
                traning = f"晚上7:00-晚上9:00在{place}参加"

            reason_text = f'\n\t现因{race_text}, 特此为以下学生请假于早{date_range}{traning}'
            new_reason_text = replace_date_in_docx(reason_text, Time)
            add_text(reason, new_reason_text, normal_size)

            # 插入表格
            add_info(cur, doc, str(c))

            # 插入日期
            EndDate = doc.add_paragraph()
            new_end_date = replace_date_in_docx(end_date, Time)
            add_text(EndDate, new_end_date, normal_size, alignment='RIGHT')

            # 插入工作室名称
            EndText = doc.add_paragraph()
            add_text(EndText, end_text, normal_size, alignment='RIGHT')
            if idx2 != len(Times) - 1:
                # 添加分页符
                doc.add_page_break()
        if idx != len(Class) - 1:
            # 添加分页符
            doc.add_page_break()

    doc.save(f"{bin.docs_path}/{','.join(str(item) for item in Class)}_{StartTime}假条.docx")

    end_conn(conn)
