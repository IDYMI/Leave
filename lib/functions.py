import ast
import datetime
import locale
from pprint import pprint
from PyQt5.QtCore import QDate
from docx import Document

import bin as bin
from lib.db.Cur import get_cur as get_cur, end_conn as end_conn
from lib.db.Select import get_info as get_info
from lib.docx.docx_add import add_info as add_info
from lib.docx.docx_add import add_text as add_text
from lib.docx.get_Times import get_Times as get_Times
from lib.docx.replace_date import replace_date_in_docx as replace_date_in_docx

# 设置默认编码为UTF-8
locale.setlocale(locale.LC_ALL, 'en_US.UTF-8')


class function(object):
    def __init__(self, ui):
        super(function, self).__init__()

        self.ui = ui
        # 点击
        self.clicks()
        # 初始化数据
        self.items()

    def clicks(self):
        self.ui.Make_Btn.clicked.connect(self.Btn_Make)  # 开始制作假条
        self.ui.Save_Btn.clicked.connect(self.Btn_Save)  # 保存假条模板
        self.ui.calendar.clicked[QDate].connect(self.date_selected)

    def items(self):
        conf = bin.conf_read()

        # 将字符串转换为列表
        self.ReasonList = ast.literal_eval(conf.get('Reason', 'reason'))
        self.PauseList = ast.literal_eval(conf.get('Pause', 'pause'))
        self.TimesList = ast.literal_eval(conf.get('Times', 'times'))
        self.ClassList = ast.literal_eval(conf.get('Class', 'class'))

        All_class = self.get_Class()
        self.ui.Class_Combox.add_items(All_class, self.ClassList)

        self.ui.ReasonBox.addItems(self.ReasonList)

        self.ui.PauseLineEdit.addItems(self.PauseList)
        self.ui.PauseLineEdit.options = self.PauseList
        self.ui.TimesLineEdit.addItems(self.TimesList)
        self.ui.TimesLineEdit.options = self.TimesList

        # 默认值
        self.ui.ReasonBox.line_edit.setText(self.ReasonList[0])
        self.ui.PauseLineEdit.line_edit.setText(self.PauseList[0])
        self.ui.TimesLineEdit.line_edit.setText(self.TimesList[0])

    def list_split(self, items, minn, maxx):
        splited_items = []  # 分页过后的结果

        index = 0  # 现在遍历到哪个同学
        # TODO 优先同导员一张表格，如果同导员人数过少小于 minn 则添加到上个导员，并且最多不超过 maxx
        for i in range(0, len(items), maxx):
            i = index if index >= i else i
            name, Class, ID, teacher, year, cnt_teacher = items[i]

            index = i  # 当前编号

            # 无前一页 或者 有前一页且满
            if (not len(splited_items)) or (len(splited_items) and len(splited_items[-1]) >= maxx):
                # 插入一页
                splited_items.append([])

            flag = 1  # 少的忽略直接添加最多添加一次
            while index + cnt_teacher < len(items):
                # 如果小于 minn, 且能添加
                if cnt_teacher < minn and flag:
                    splited_items[-1] += items[index: index + cnt_teacher]
                    index += cnt_teacher
                    cnt_teacher = items[index][5]
                    if len(splited_items[-1]) + cnt_teacher > maxx:
                        flag = 0
                # 超过 maxx， 且能添加
                elif cnt_teacher > maxx:
                    cnt = 0
                    while cnt < cnt_teacher // maxx:
                        if len(splited_items[-1]) != 0:
                            splited_items.append(items[index: index + maxx])
                        else:
                            splited_items[-1] += (items[index: index + maxx])
                        index += maxx
                        cnt += 1
                    if cnt_teacher % maxx > 0:
                        splited_items.append(
                            items[index: index + cnt_teacher % maxx])
                        index += cnt_teacher % maxx
                        cnt_teacher = items[index][5]
                # 不小于 minn， 不大于 maxx， 当前页能添加完
                elif len(splited_items[-1]) + cnt_teacher <= maxx:
                    splited_items[-1] += items[index: index + cnt_teacher]
                    index += cnt_teacher
                    cnt_teacher = items[index][5]
                    # if len(splited_items[-1]) == maxx:
                    #     splited_items.append([])
                # 不小于 minn， 不大于 maxx， 当前页不能添加完
                elif len(splited_items[-1]) + cnt_teacher > maxx:
                    remain = maxx - len(splited_items[-1])
                    splited_items[-1] += items[index: index + remain]
                    splited_items.append(
                        items[index + remain: index + cnt_teacher + 1])
                    index += cnt_teacher
                    cnt_teacher = items[index][5]
                else:
                    break
        # 如果有剩余
        if index != len(items):
            splited_items[-1] += items[index: index + cnt_teacher]
        # pprint(splited_items)
        return splited_items

    def Btn_Make(self):

        self.ui.Make_Btn.setEnabled(False)

        title_text = '请假条'  # 标题

        StartTime = self.ui.StartTimeLineEdit.text()  # 获取开始日期
        StartTime = datetime.datetime.now().strftime(
            "%m月%d日") if '点击选择开始日期' in StartTime else StartTime  # 判断如果没有赋值, 开始时间为今天

        race_text = self.ui.ReasonBox.line_edit.text()  # 请假理由
        Pause = int(self.ui.PauseLineEdit.line_edit.text())  # 一次时长
        T = int(self.ui.TimesLineEdit.line_edit.text())  # 连续几次

        Class = self.ui.Class_Combox.getCheckItems()  # 选中的届别

        year = datetime.datetime.now().year  # 当前年份
        traning = "1"  # 默认地点
        normal_size = 12  # 字体的默认大小
        end_date = f'{year}年01月01日'  # 默认批准时间
        date_range = '01月01日 - 01月02日'  # 默认请假时间
        end_text = '计算机与软件学院计科ACM集训队'  # 落款

        Times = get_Times(StartTime, Times=T, Pause=Pause)  # 计算得到需要的日期列表

        conn, cur = get_cur()  # 查询数据库

        doc = Document()  # 新建一个空白页

        for idx3, c in enumerate(Class):
            """遍历届数"""
            data = get_info(cur, str(c))
            # print(data)
            data_split = self.list_split(data, 3, 8)  # 按照每页不少于3，不多于8

            for idx, data_splited in enumerate(data_split):
                for idx2, Time in enumerate(Times):
                    """遍历时间域"""

                    # 插入标题
                    title = doc.add_heading()
                    add_text(title, title_text, 26, Bold=True, alignment='CENTER')

                    # 插入原因
                    reason = doc.add_paragraph()

                    if traning != '全天':
                        traning = f"早7：00-晚10：00"
                    else:
                        traning = f"晚上7:00-晚上9:00"

                    reason_text = f'\n\t现因{race_text}, 特此为以下学生请假于早{date_range}{traning}'
                    new_reason_text = replace_date_in_docx(reason_text, Time)
                    add_text(reason, new_reason_text, normal_size)

                    # 插入表格
                    add_info(cur, doc, data_splited)

                    # 插入日期
                    EndDate = doc.add_paragraph()
                    new_end_date = replace_date_in_docx(end_date, Time)
                    add_text(EndDate, new_end_date, normal_size, alignment='RIGHT')

                    # 插入工作室名称
                    EndText = doc.add_paragraph()
                    add_text(EndText, end_text, normal_size, alignment='RIGHT')

                    if idx2 != len(Times) - 1:
                        # 添加分页符
                        doc.add_page_break()
                if idx != len(data_split) - 1  :
                    # 添加分页符
                    doc.add_page_break()
            if idx3 != len(Class) - 1:
                # 添加分页符
                doc.add_page_break()

        save_place = f"{bin.docs_path}/{','.join(str(item) for item in Class)}_{StartTime}假条.docx"
        doc.save(save_place)

        end_conn(conn)

        self.ui.send_message(save_place)

        self.ui.Make_Btn.setEnabled(True)

    def move_to_start(self, lis, elem):
        lis.insert(0, lis.pop(lis.index(elem)))

    def Btn_Save(self):
        race_text = self.ui.ReasonBox.line_edit.text()  # 请假理由
        Pause = self.ui.PauseLineEdit.line_edit.text()  # 一次时长
        T = self.ui.TimesLineEdit.line_edit.text()  # 连续几次
        C = self.ui.Class_Combox.qLineEdit.text().split(',') if self.ui.Class_Combox.qLineEdit.text() else [] # 添加的班级

        if race_text != self.ReasonList[0] or Pause != self.PauseList[0] or T != self.TimesList[0] or C != self.ClassList:
            # 如果发生改变, 调换顺序
            self.move_to_start(self.ReasonList, race_text)
            self.move_to_start(self.PauseList, Pause)
            self.move_to_start(self.TimesList, T)
            self.ClassList = C

            Data = {
                "Reason": self.ReasonList,
                "Pause": self.PauseList,
                "Times": self.TimesList,
                "Class": self.ClassList
            }

            # 更新配置
            bin.conf_save(Data)

    def date_selected(self, date):
        self.ui.StartTimeLineEdit.setText(date.toString("MM月dd日"))
        # self.hide_calendar()

    def setTextView(self, item, text):
        item.setToolTip(text)
        item.setPlaceholderText(text)

    def get_Class(self):
        conn, cur = get_cur()
        options = get_info(cur, misson=3)
        end_conn(conn)
        return options
